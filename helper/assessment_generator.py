import json
from helper.llm import get_completion
from docx import Document

def generate_assessment_with_answers(jd_matched_skills):
    """
    Uses LLM to generate a 5-question assessment with answers in JSON format.
    Each question includes the skill it assesses.

    Args:
        jd_matched_skills (list): List of skills relevant to the job description.

    Returns:
        list: List of dictionaries, each containing a question, answer, and skill.
    """
    prompt = f"""
    Based on the following skills required for the job, create an assessment with questions in JSON format. 
    Each question should assess the candidate's ability in each skill. Provide a brief answer for each question as part of an answer key, 
    and include the skill being assessed for each question.

    Skills Required:
    {', '.join([skill['Skill'] for skill in jd_matched_skills])}

    Output the result in JSON format, with each object containing a "question" field, an "answer" field, and a "skill" field. Here is the format:
    [
      {{"question": "Question text 1", "answer": "Answer text 1", "skill": "Skill being assessed 1"}},
      {{"question": "Question text 2", "answer": "Answer text 2", "skill": "Skill being assessed 2"}},
      ...
    ]
    """
    
    # Generate structured assessment from LLM
    response = get_completion(prompt)
     # Clean up the response to remove any extra formatting like ```json
    response = response.strip("```json").strip("```").strip()
    
    assessment_data = json.loads(response)
    
    return assessment_data

def create_candidate_docs(candidate_names, assessment_data):
    """
    Creates a Word document for each candidate with the assessment questions only.

    Args:
        candidate_names (list): List of candidate names.
        assessment_data (list): List of dictionaries, each containing a question and answer.
    """
    for candidate_name in candidate_names:
        doc = Document()
        doc.add_heading(f"Assessment for {candidate_name}", level=1)
        doc.add_paragraph("Please answer the following questions to the best of your ability.\n")
        
        for item in assessment_data:
            doc.add_paragraph(item["question"], style='List Number')

        filename = f"{candidate_name.replace(' ', '_')}_Assessment.docx"
        doc.save(filename)

def create_answer_key_doc(assessment_data):
    """
    Creates a Word document with the answer key for the assessment.

    Args:
        assessment_data (list): List of dictionaries, each containing a question, answer, and skill.
    """
    doc = Document()
    doc.add_heading("Assessment Answer Key", level=1)

    for i, item in enumerate(assessment_data, start=1):
        doc.add_paragraph(f"Q{i}. {item['question']}", style='List Number')
        doc.add_paragraph(f"Skill: {item['skill']}", style='List Bullet')  # Include skill being assessed
        doc.add_paragraph(f"Answer: {item['answer']}", style='List Bullet')
    
    doc.save("Assessment_Answer_Key.docx")

