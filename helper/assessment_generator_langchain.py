import streamlit as st
from langchain.chains import SequentialChain, LLMChain
from langchain.prompts import PromptTemplate
from langchain.llms import OpenAI
from docx import Document
import json

# Initialize the LangChain LLM
llm = OpenAI(api_key=st.secrets["PERSONAL_OPENAI_API_KEY"])

# Define prompt templates
question_generation_prompt = PromptTemplate(
    input_variables=["jd_matched_skills"],
    template="""
    Based on the following skills required for the job, create an assessment with 5 questions. 
    Each question should assess the candidate's ability in one of the skills. 

    Skills Required:
    {jd_matched_skills}

    Output the result in JSON format as follows:
    [
      {{"question": "Question text 1", "skill": "Skill being assessed 1"}},
      {{"question": "Question text 2", "skill": "Skill being assessed 2"}},
      ...
    ]
    """
)

answer_generation_prompt = PromptTemplate(
    input_variables=["questions"],
    template="""
    Based on the questions below, provide a brief answer for each question as part of an answer key. 
    Output the result in JSON format as follows:
    [
      {{"question": "Question text 1", "answer": "Answer text 1", "skill": "Skill being assessed 1"}},
      {{"question": "Question text 2", "answer": "Answer text 2", "skill": "Skill being assessed 2"}},
      ...
    ]

    Questions:
    {questions}
    """
)

# Create individual LLMChains for each prompt
question_chain = LLMChain(llm=llm, prompt=question_generation_prompt, output_key="questions")
answer_chain = LLMChain(llm=llm, prompt=answer_generation_prompt, output_key="assessment_data")

# Combine them into a SequentialChain
sequential_chain = SequentialChain(
    chains=[question_chain, answer_chain],
    input_variables=["jd_matched_skills"],
    output_variables=["assessment_data"]
)

def generate_assessment_with_answers(jd_matched_skills):
    """
    Uses LangChain to generate a 5-question assessment with answers in JSON format.
    Each question includes the skill it assesses.

    Args:
        jd_matched_skills (list): List of skills relevant to the job description.

    Returns:
        list: List of dictionaries, each containing a question, answer, and skill.
    """
    # Convert jd_matched_skills to a formatted string for the prompt
    skills_text = ', '.join([skill['Skill'] for skill in jd_matched_skills])

    # Run the sequential chain with skills_text as input
    chain_output = sequential_chain({"jd_matched_skills": skills_text})

    # Retrieve and parse the assessment data
    assessment_data_response = chain_output["assessment_data"]
    assessment_data_response = assessment_data_response.strip("```json").strip("```").strip()
    assessment_data = json.loads(assessment_data_response)

    return assessment_data

def create_candidate_docs(candidate_names, assessment_data):
    """
    Creates a Word document for each candidate with the assessment questions only.

    Args:
        candidate_names (list): List of candidate names.
        assessment_data (list): List of dictionaries, each containing a question and answer.
    """
    for candidate_name in candidate_names:
        doc = Document()
        doc.add_heading(f"Assessment for {candidate_name}", level=1)
        doc.add_paragraph("Please answer the following questions to the best of your ability.\n")
        
        for item in assessment_data:
            doc.add_paragraph(item["question"], style='List Number')

        filename = f"{candidate_name.replace(' ', '_')}_Assessment.docx"
        doc.save(filename)

def create_answer_key_doc(assessment_data):
    """
    Creates a Word document with the answer key for the assessment.

    Args:
        assessment_data (list): List of dictionaries, each containing a question, answer, and skill.
    """
    doc = Document()
    doc.add_heading("Assessment Answer Key", level=1)

    for i, item in enumerate(assessment_data, start=1):
        doc.add_paragraph(f"Q{i}. {item['question']}", style='List Number')
        doc.add_paragraph(f"Skill: {item['skill']}", style='List Bullet')  # Include skill being assessed
        doc.add_paragraph(f"Answer: {item['answer']}", style='List Bullet')
    
    doc.save("Assessment_Answer_Key.docx")
